import os
import io
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import chromadb
from sentence_transformers import SentenceTransformer
import streamlit as st

load_dotenv()
groq_key = os.environ.get("GROQ_API_KEY")
groq_client = Groq(api_key=groq_key)

# ── Read requirements from an in-memory file (uploaded bytes) ─
def read_requirements_from_bytes(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    requirements = []
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_paragraphs.append(para)
    for paragraph in all_paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text and not style_name.startswith("Heading"):
            requirements.append(text)
    return requirements

# ── Compress requirements to reduce token usage ───────────────
def compress_requirements(requirements):
    compressed = []
    for req in requirements:
        r = req
        filler = [
            "The system shall ", "The system should ", "The system will ",
            "The system must ", "The system is required to ",
            "It is required that the system ", "The application shall ",
            "Users shall be able to ", "Users should be able to ",
            "The platform shall ", "The solution shall ",
        ]
        for f in filler:
            r = r.replace(f, "")
        if r:
            r = r[0].upper() + r[1:]
        skip_keywords = ["Note:", "Note from", "TBD", "To be confirmed", "?"]
        if any(r.startswith(k) for k in skip_keywords):
            continue
        if len(r) < 20:
            continue
        compressed.append(r.strip())
    seen = set()
    unique = []
    for r in compressed:
        if r not in seen:
            seen.add(r)
            unique.append(r)
    return unique

# ── Build ChromaDB vector database ───────────────────────────
def build_vector_db(requirements):
    model = SentenceTransformer("all-MiniLM-L6-v2")
    chroma_client = chromadb.Client()
    collection = chroma_client.create_collection("requirements")
    for i, req in enumerate(requirements):
        embedding = model.encode(req).tolist()
        collection.add(
            documents=[req],
            embeddings=[embedding],
            ids=[f"req_{i}"]
        )
    return collection, model

# ── Search requirements by meaning ───────────────────────────
def search_requirements(question, collection, model, n=3):
    question_embedding = model.encode(question).tolist()
    results = collection.query(
        query_embeddings=[question_embedding],
        n_results=min(n, len(st.session_state.requirements))
    )
    return results["documents"][0]

# ── Ask Groq AI ───────────────────────────────────────────────
def ask_groq(prompt):
    response = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# ── Generate user stories ─────────────────────────────────────
def generate_user_stories(requirements):
    numbered = "\n".join([f"{i+1}. {r}" for i, r in enumerate(requirements)])
    prompt = f"""
You are a senior Business Analyst. Based on the requirements below, generate user stories in this exact format:

As a [user type], I want to [action], so that [benefit].

Acceptance Criteria:
- [criterion 1]
- [criterion 2]
- [criterion 3]

Generate one user story for each requirement. Be specific and professional.

Requirements:
{numbered}
"""
    return ask_groq(prompt)

# ── Run gap analysis ──────────────────────────────────────────
def run_gap_analysis(requirements):
    numbered = "\n".join([f"{i+1}. {r}" for i, r in enumerate(requirements)])
    prompt = f"""
You are a senior Business Analyst and QA expert performing a requirements review.

Analyze each requirement below and identify any of these gap types:
- VAGUE: The requirement is ambiguous or unclear
- MISSING: Something important is not defined
- UNTESTABLE: The requirement cannot be objectively verified
- CONFLICT: The requirement contradicts another requirement

For each gap found, respond in this exact format:

Requirement #: [number]
Requirement: [the original text]
Gap Type: [VAGUE / MISSING / UNTESTABLE / CONFLICT]
Severity: [HIGH / MEDIUM / LOW]
Issue: [one sentence describing the problem]
Recommendation: [one sentence on how to fix it]

---

If a requirement has no gaps, write:
Requirement #: [number]
Status: PASS

---

Requirements:
{numbered}
"""
    return ask_groq(prompt)

# ── Build downloadable .docx from AI output ───────────────────
def build_docx(title, content):
    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:]).font.size = Pt(10)
        elif line == "---":
            doc.add_paragraph("─" * 60)
        else:
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(10)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ── Streamlit UI ──────────────────────────────────────────────
st.set_page_config(page_title="AI Backlog Intelligence", page_icon="🤖", layout="wide")
st.title("🤖 AI Requirements & Backlog Intelligence System")
st.caption("Upload any requirements document to generate user stories, detect gaps, and chat with your backlog")

# ── Session state initialization ─────────────────────────────
if "requirements" not in st.session_state:
    st.session_state.requirements = []
if "collection" not in st.session_state:
    st.session_state.collection = None
if "embed_model" not in st.session_state:
    st.session_state.embed_model = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "doc_name" not in st.session_state:
    st.session_state.doc_name = None

# ── Sidebar — file upload ─────────────────────────────────────
with st.sidebar:
    st.header("📄 Upload Document")
    uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

    if uploaded_file:
        if uploaded_file.name != st.session_state.doc_name:
            with st.spinner("Processing document..."):
                file_bytes = uploaded_file.read()
                requirements = read_requirements_from_bytes(file_bytes)
                requirements = compress_requirements(requirements)
                collection, embed_model = build_vector_db(requirements)
                st.session_state.requirements = requirements
                st.session_state.collection = collection
                st.session_state.embed_model = embed_model
                st.session_state.messages = []
                st.session_state.doc_name = uploaded_file.name
            st.success(f"Loaded {len(requirements)} requirements")

    if st.session_state.requirements:
        st.divider()
        st.header("📋 Requirements")
        for i, req in enumerate(st.session_state.requirements, 1):
            st.write(f"{i}. {req}")

# ── Main area — tabs ──────────────────────────────────────────
if not st.session_state.requirements:
    st.info("👈 Upload a .docx requirements document from the sidebar to get started")
    st.stop()

tab1, tab2, tab3 = st.tabs(["💬 Chat", "📝 User Stories", "🔍 Gap Analysis"])

# ── TAB 1: Chat ───────────────────────────────────────────────
with tab1:
    st.subheader("Chat with your requirements")

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if question := st.chat_input("Ask something about your requirements..."):
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                relevant = search_requirements(
                    question,
                    st.session_state.collection,
                    st.session_state.embed_model
                )
                context = "\n".join([f"- {r}" for r in relevant])
                prompt = f"""
You are an AI assistant helping a Business Analyst understand their requirements document.
Use only the requirements below to answer the question. Be specific and professional.

Requirements:
{context}

Question: {question}
"""
                answer = ask_groq(prompt)
                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})

# ── TAB 2: User Stories ───────────────────────────────────────
with tab2:
    st.subheader("Generate User Stories")
    st.caption("AI will generate one INVEST-compliant user story per requirement")

    if "gap_analysis" not in st.session_state:
        st.warning("⚠️ Recommended: Run Gap Analysis first to identify and fix requirement issues before generating user stories.")

    if st.button("Generate User Stories", type="primary"):
        with st.spinner("Generating user stories..."):
            result = generate_user_stories(st.session_state.requirements)
            st.session_state.user_stories = result

    if "user_stories" in st.session_state:
        st.markdown(st.session_state.user_stories)
        st.divider()
        docx_buffer = build_docx("User Stories", st.session_state.user_stories)
        st.download_button(
            label="⬇️ Download as .docx",
            data=docx_buffer,
            file_name="user_stories.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ── TAB 3: Gap Analysis ───────────────────────────────────────
with tab3:
    st.subheader("Requirements Gap Analysis")
    st.caption("AI will flag vague, missing, untestable, or conflicting requirements")

    if st.button("Run Gap Analysis", type="primary"):
        with st.spinner("Analyzing requirements..."):
            result = run_gap_analysis(st.session_state.requirements)
            st.session_state.gap_analysis = result

    if "gap_analysis" in st.session_state:
        st.markdown(st.session_state.gap_analysis)
        st.divider()
        docx_buffer = build_docx("Requirements Gap Analysis", st.session_state.gap_analysis)
        st.download_button(
            label="⬇️ Download as .docx",
            data=docx_buffer,
            file_name="gap_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )